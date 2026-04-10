import docx
import os

folder = r"c:\Users\pares\Documents\covenant"
files = ["Balance Sheet - Lunar.docx", "Cash Flow Statements - Lunar.docx", "Income Statement - Lunar.docx", "Mock Credit Agreement.docx"]

with open(os.path.join(folder, "output_data.txt"), "w", encoding="utf-8") as f:
    for file in files:
        f.write(f"\n--- {file} ---\n")
        try:
            doc = docx.Document(os.path.join(folder, file))
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text.strip() + "\n")
            for table in doc.tables:
                for row in table.rows:
                    f.write("\t".join([cell.text.strip() for cell in row.cells]) + "\n")
        except Exception as e:
            f.write(f"Error reading {file}: {e}\n")
